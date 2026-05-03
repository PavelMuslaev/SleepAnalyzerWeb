from docx import Document
from docx.shared import Inches, Pt, RGBColor  # добавлен RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import List, Tuple, Optional


class DocxGenerator:
    @staticmethod
    def generate(analysis_dict: dict,
                 num_records: int = 2,
                 images: Optional[List[Tuple[str, bytes]]] = None):
        """Генерирует документ Word с отчётом. images – список (подпись, PNG-байты)."""
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Заголовок
        title = doc.add_heading('Отчёт SleepAnalyzer', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Информация о пользователе
        if analysis_dict.get('user_age'):
            doc.add_paragraph(
                f"Возраст: {analysis_dict['user_age']} лет   Пол: {analysis_dict['user_gender'] or 'не указан'}"
            )

        # Примечание при 1 записи
        if num_records <= 2:
            p = doc.add_paragraph()
            run = p.add_run("⚠️ Внимание: у вас меньше трех записей сна. Для более точного анализа и отслеживания трендов рекомендуется загрузить данные за 3–7 дней.")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 80, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()

        # Основные показатели
        doc.add_heading('Основные показатели', level=2)
        table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Значение'

        def add_row(label, value, fmt='.1f'):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            if isinstance(value, float):
                row_cells[1].text = format(value, fmt)
            else:
                row_cells[1].text = str(value) if value is not None else '—'

        add_row('Общая длительность сна (ч)', analysis_dict['avg_total_sleep']/60)
        add_row('Время в постели (ч)', analysis_dict['avg_time_in_bed']/60)
        add_row('Эффективность сна (%)', analysis_dict['avg_efficiency'])
        add_row('Глубокий сон (мин)', analysis_dict['avg_deep_sleep'])
        add_row('Лёгкий сон (мин)', analysis_dict['avg_light_sleep'])
        add_row('REM-сон (мин)', analysis_dict['avg_rem_sleep'])
        add_row('Латентность (мин)', analysis_dict['avg_latency'], '.0f')
        add_row('WASO (мин)', analysis_dict['avg_waso'], '.0f')
        add_row('Фрагментация (проб/ч)', analysis_dict['avg_fragmentation'])
        add_row('Регулярность (стд. откл, мин)', analysis_dict['sleep_regularity'])
        add_row('Консистентность времени сна (стд. откл, мин)', analysis_dict['bedtime_consistency'])
        add_row('Средний пульс', analysis_dict['avg_heart_rate'], '.0f')
        add_row('Средняя HRV', analysis_dict['avg_hrv'], '.0f')
        add_row('Частота дыхания', analysis_dict['avg_respiratory_rate'], '.1f')
        add_row('Движения', analysis_dict['avg_movement'], '.1f')
        add_row('Субъективная оценка', analysis_dict['avg_subjective_rating'], '.1f')
        add_row('Sleep Score (из 100)', analysis_dict['sleep_score'])
        add_row('Долг сна (ч)', analysis_dict['sleep_debt_minutes']/60)
        add_row('Тренд', analysis_dict['trend'])

        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        for rec in analysis_dict.get('recommendations', []):
            p = doc.add_paragraph()
            level = rec.get('level', 'info')
            if level == 'danger':
                run = p.add_run('⚠️ (Важно) ')
                run.bold = True
                p.add_run(rec['message'])
            elif level == 'warning':
                run = p.add_run('⚡ (Обратить внимание) ')
                run.bold = True
                p.add_run(rec['message'])
            else:
                p.add_run(rec['message'])

        # Графики
        if images:
            doc.add_heading('Графики', level=2)
            for caption, img_bytes in images:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                run.bold = True
                # Вставляем изображение с уменьшением до 5.5 дюймов
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf